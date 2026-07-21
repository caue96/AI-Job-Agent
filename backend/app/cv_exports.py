"""Deterministic document rendering and traversal-safe private export storage."""

from __future__ import annotations

import hashlib
import io
import re
import uuid
from html import escape
from pathlib import Path

from app.cover_letter_schemas import CoverLetterContent
from app.cv_schemas import CvProfileDraft, CvValue


class CvExportError(ValueError):
    pass


def _value(item: CvValue) -> str:
    return "" if item.value is None else str(item.value).strip()


def _sections(profile: CvProfileDraft) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    summary = _value(profile.professional_summary)
    if summary:
        sections.append(("Professional summary", [summary]))
    skills = [item.value for item in profile.technical_skills]
    if skills:
        sections.append(("Technical skills", [", ".join(skills)]))
    for entry in profile.employment:
        heading = " | ".join(filter(None, (_value(entry.title), _value(entry.company))))
        dates = " - ".join(filter(None, (_value(entry.start_date), _value(entry.end_date))))
        lines = [dates] if dates else []
        lines.extend(item.value for item in entry.achievements)
        lines.extend(item.value for item in entry.responsibilities)
        if heading:
            sections.append((heading, lines))
    if profile.projects:
        lines = []
        for project in profile.projects:
            label = _value(project.name)
            description = _value(project.description)
            lines.append(": ".join(filter(None, (label, description))))
        sections.append(("Projects", [line for line in lines if line]))
    education = [
        " | ".join(
            filter(
                None,
                (_value(item.qualification), _value(item.field_of_study), _value(item.institution)),
            )
        )
        for item in profile.education
    ]
    if any(education):
        sections.append(("Education", [item for item in education if item]))
    certifications = [
        " | ".join(filter(None, (_value(item.name), _value(item.issuer))))
        for item in profile.certifications
    ]
    if any(certifications):
        sections.append(("Certifications", [item for item in certifications if item]))
    languages = [item.value for item in profile.languages]
    if languages:
        sections.append(("Languages", [", ".join(languages)]))
    return sections


def render_docx(profile: CvProfileDraft) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.65)
    section.left_margin = section.right_margin = Inches(0.75)
    name = _value(profile.personal.full_name) or "Candidate"
    document.add_heading(name, 0)
    headline = _value(profile.headline)
    if headline:
        document.add_paragraph(headline, style="Subtitle")
    contact = [
        _value(profile.personal.city),
        _value(profile.personal.country),
        _value(profile.personal.email),
        _value(profile.personal.phone),
        _value(profile.personal.linkedin_url),
    ]
    contact = [item for item in contact if item]
    if contact:
        document.add_paragraph(" | ".join(contact))
    for heading, lines in _sections(profile):
        document.add_heading(heading, level=1)
        for line in lines:
            document.add_paragraph(line, style="List Bullet" if len(lines) > 1 else None)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(profile: CvProfileDraft) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=_value(profile.personal.full_name) or "CV",
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("CvTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=18)
    subtitle = ParagraphStyle(
        "CvSubtitle", parent=styles["Normal"], alignment=TA_CENTER, fontSize=10
    )
    story = [Paragraph(escape(_value(profile.personal.full_name) or "Candidate"), title)]
    headline = _value(profile.headline)
    if headline:
        story.append(Paragraph(escape(headline), subtitle))
    contact = [
        _value(profile.personal.city),
        _value(profile.personal.country),
        _value(profile.personal.email),
        _value(profile.personal.phone),
        _value(profile.personal.linkedin_url),
    ]
    story.append(Paragraph(escape(" | ".join(item for item in contact if item)), subtitle))
    story.append(Spacer(1, 5 * mm))
    for heading, lines in _sections(profile):
        story.append(Paragraph(escape(heading), styles["Heading2"]))
        for line in lines:
            story.append(Paragraph(f"&#8226; {escape(line)}", styles["BodyText"]))
        story.append(Spacer(1, 2 * mm))
    document.build(story)
    return output.getvalue()


class LocalCvExportStorage:
    def __init__(self, configured_path: str):
        self.root = Path(configured_path).expanduser().resolve()
        self.root.mkdir(parents=True, exist_ok=True)

    def _path(self, key: str) -> Path:
        if not re.fullmatch(r"[0-9a-f]{32}\.(?:pdf|docx|txt)", key):
            raise CvExportError("Invalid CV export identifier")
        target = (self.root / key).resolve()
        if target.parent != self.root:
            raise CvExportError("Invalid CV export path")
        return target

    def store(self, format_name: str, content: bytes) -> tuple[str, str, int]:
        if format_name not in {"pdf", "docx", "txt"}:
            raise CvExportError("Unsupported CV export format")
        key = f"{uuid.uuid4().hex}.{format_name}"
        self._path(key).write_bytes(content)
        return key, hashlib.sha256(content).hexdigest(), len(content)

    def path_for(self, key: str) -> Path:
        path = self._path(key)
        if not path.is_file():
            raise CvExportError("CV export file is unavailable")
        return path

    def delete(self, key: str) -> None:
        self._path(key).unlink(missing_ok=True)


def render_export(profile: CvProfileDraft, format_name: str) -> bytes:
    return render_pdf(profile) if format_name == "pdf" else render_docx(profile)


def render_cover_letter_txt(content: CoverLetterContent) -> bytes:
    values = [
        content.candidate_name,
        content.contact_line,
        content.date,
        content.company,
        content.job_title,
        "",
        content.greeting,
        "",
        *(paragraph.text + "\n" for paragraph in content.paragraphs),
        content.signoff,
        content.candidate_name,
    ]
    return "\n".join(value for value in values if value is not None).encode("utf-8")


def render_cover_letter_docx(content: CoverLetterContent) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.85)
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(content.candidate_name)
    run.bold = True
    run.font.size = Pt(16)
    if content.contact_line:
        contact = document.add_paragraph(content.contact_line)
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(content.date)
    document.add_paragraph(content.company)
    document.add_paragraph(content.job_title)
    document.add_paragraph(content.greeting)
    for paragraph in content.paragraphs:
        document.add_paragraph(paragraph.text)
    document.add_paragraph(content.signoff)
    document.add_paragraph(content.candidate_name)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(11)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_cover_letter_pdf(content: CoverLetterContent) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"Cover letter - {content.job_title}",
        author=content.candidate_name,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "CoverLetterTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=16
    )
    contact = ParagraphStyle(
        "CoverLetterContact", parent=styles["Normal"], alignment=TA_CENTER, fontSize=9
    )
    body = ParagraphStyle(
        "CoverLetterBody",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )
    story = [Paragraph(escape(content.candidate_name), title)]
    if content.contact_line:
        story.append(Paragraph(escape(content.contact_line), contact))
    story.extend(
        [
            Spacer(1, 6 * mm),
            Paragraph(escape(content.date), body),
            Paragraph(escape(content.company), body),
            Paragraph(escape(content.job_title), body),
            Spacer(1, 3 * mm),
            Paragraph(escape(content.greeting), body),
        ]
    )
    story.extend(Paragraph(escape(paragraph.text), body) for paragraph in content.paragraphs)
    story.extend(
        [
            Spacer(1, 2 * mm),
            Paragraph(escape(content.signoff), body),
            Paragraph(escape(content.candidate_name), body),
        ]
    )
    document.build(story)
    return output.getvalue()


def render_cover_letter_export(content: CoverLetterContent, format_name: str) -> bytes:
    if format_name == "txt":
        return render_cover_letter_txt(content)
    if format_name == "docx":
        return render_cover_letter_docx(content)
    if format_name == "pdf":
        return render_cover_letter_pdf(content)
    raise CvExportError("Unsupported cover-letter export format")
